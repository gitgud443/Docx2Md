#!/usr/bin/env python3

import os
import sys
import docx
import unicodedata
from docx.oxml.text.run import CT_R
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
import re

def is_valid_unicode(char):
    """Check if a character is a valid Unicode character."""
    try:
        unicodedata.name(char)
        return True
    except ValueError:
        return False

def analyze_docx(file_path):
    """Analyze a DOCX file for non-Unicode characters."""
    print(f"\nAnalyzing: {os.path.basename(file_path)}")
    
    # Load the document
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        print(f"Error opening document: {e}")
        return
    
    # Track non-Unicode characters
    non_unicode_chars = {}
    total_chars = 0
    
    # Function to process text and find non-Unicode characters
    def process_text(text, location):
        nonlocal total_chars
        total_chars += len(text)
        
        for i, char in enumerate(text):
            if char != "\n" and char != "\t":
                if not is_valid_unicode(char):
                    char_repr = repr(char)[1:-1]  # Get string representation without quotes
                    if char_repr not in non_unicode_chars:
                        non_unicode_chars[char_repr] = []
                    non_unicode_chars[char_repr].append(f"{location} (position {i})")
    
    # Process paragraphs
    for i, para in enumerate(doc.paragraphs):
        process_text(para.text, f"Paragraph {i+1}")
    
    # Process tables
    for i, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    process_text(para.text, f"Table {i+1}, Row {row_idx+1}, Column {col_idx+1}, Paragraph {para_idx+1}")
    
    # Process runs directly (might catch more characters)
    for i, para in enumerate(doc.paragraphs):
        for j, run in enumerate(para.runs):
            process_text(run.text, f"Paragraph {i+1}, Run {j+1}")
    
    # Report findings
    if non_unicode_chars:
        print(f"Found {len(non_unicode_chars)} types of non-Unicode characters:")
        for char, locations in non_unicode_chars.items():
            print(f"  - Character: '{char}' (hex: {' '.join(hex(ord(c))[2:] for c in char if ord(c) < 0x110000)})")
            print(f"    Appears {len(locations)} times")
            # Print first 5 locations
            for loc in locations[:5]:
                print(f"    - {loc}")
            if len(locations) > 5:
                print(f"    - ... and {len(locations) - 5} more locations")
    else:
        print("No non-Unicode characters found.")
    
    print(f"Total characters analyzed: {total_chars}")
    
    return non_unicode_chars

def analyze_directory(directory):
    """Analyze all DOCX files in a directory."""
    docx_files = [f for f in os.listdir(directory) if f.endswith('.docx')]
    
    if not docx_files:
        print(f"No DOCX files found in {directory}")
        return
    
    print(f"Found {len(docx_files)} DOCX files to analyze")
    
    # Summary of findings
    summary = {}
    
    for file in docx_files:
        file_path = os.path.join(directory, file)
        non_unicode = analyze_docx(file_path)
        
        if non_unicode:
            summary[file] = len(non_unicode)
    
    # Print overall summary
    print("\n=== SUMMARY ===")
    if summary:
        print(f"{len(summary)} files contain non-Unicode characters:")
        for file, count in sorted(summary.items(), key=lambda x: x[1], reverse=True):
            print(f"  - {file}: {count} types of non-Unicode characters")
    else:
        print("No non-Unicode characters found in any files.")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python detect_non_unicode.py <file.docx or directory>")
        sys.exit(1)
    
    path = sys.argv[1]
    
    if os.path.isdir(path):
        analyze_directory(path)
    elif os.path.isfile(path) and path.endswith('.docx'):
        analyze_docx(path)
    else:
        print("Please provide a valid DOCX file or directory containing DOCX files.")
        sys.exit(1)
