import sys
import json
import shutil
import os
from docx import Document
from docx.oxml.ns import qn

def has_code_font(para):
    code_fonts = ['courier new', 'helvetica 45 light', 'consolas', 'Courier New', 'Consolas', 'source code pro', 'source code']
    for run in para.runs:
        if run.font.name and run.font.name.lower() in code_fonts:
            return True
    return False

def has_shading(para):
    pPr = para._p.pPr
    if pPr is not None:
        shd = pPr.find(qn('w:shd'))
        if shd is not None:
            fill = shd.get(qn('w:fill'))
            if fill and fill != 'auto':
                return True
    return False

def has_frame(para):
    pPr = para._p.pPr
    if pPr is not None:
        framePr = pPr.find(qn('w:framePr'))
        if framePr is not None:
            return True
    return False

def has_border(para):
    pPr = para._p.pPr
    if pPr is not None:
        borders = pPr.find(qn('w:pBdr'))
        if borders is not None:
            return True
    return False

def is_code_paragraph(para):
    # Exclude titles like table of contents
    excluded_titles = {
        'table of contents',
        'contents',
        'table Of contents',
        'table Of Contents',
        'Table of contents',
        'Table Of Contents',
        'Table of content',
        'table of content',
        'Table des matières',
        'Contents',
        'contents'
    }

    if para.text.strip().lower() in {title.lower() for title in excluded_titles}:
        return False

    # Classic code block detection
    if para.style and para.style.name.lower() == 'code':
        return True
    if has_code_font(para):
        return True
    if para.style and para.style.name in ['CODE', 'Code0', 'StyleCourierNew9ptLeft029cmRight048cmBox', 'cadre', 'Config',  'Configuration 1',]:
        return True
    if has_shading(para):
        return True
    if has_border(para) or has_frame(para):
        return True
    return False

def replace_code_blocks_by_markers(doc):
    code_blocks = []
    current_block = []
    in_block = False
    block_id = 0

    # First step : extracting the code blocks
    for para in doc.paragraphs:
        if is_code_paragraph(para):
            current_block.append(para.text)
            in_block = True
        else:
            if in_block:
                block_id += 1
                code_blocks.append('\n'.join(current_block))
                current_block = []
                in_block = False

    if in_block:
        block_id += 1
        code_blocks.append('\n'.join(current_block))

    # Second step : replace only the first paragraph of each code block with the marker, others are emptied
    # And empty the contents of the other paragraphs
    block_id = 0
    in_block = False
    first_para_in_block = False

    for para in doc.paragraphs:
        if is_code_paragraph(para):
            if not in_block:
                block_id += 1
                in_block = True
                first_para_in_block = True
            if first_para_in_block:
                # Replace the first paragraph's text with the marker
                for run in para.runs:
                    run.text = ''
                para.add_run(f'@@CODEBLOCK_{block_id}@@')
                first_para_in_block = False
            else:
                # Empty the contents of the other paragraphs in the block (to avoid clones)
                for run in para.runs:
                    run.text = ''
        else:
            if in_block:
                in_block = False

    return code_blocks

def main(input_docx_path, output_docx_path, output_json_path):
    # Create the output directory if not exists
    os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
    # Copy the source file to its destination
    shutil.copy2(input_docx_path, output_docx_path)
    # Open the copied file for modification
    doc = Document(output_docx_path)
    # Modifying the file (replacing code blocks with markers)
    code_blocks = replace_code_blocks_by_markers(doc)
    # Saving the modifications inside the copied file
    doc.save(output_docx_path)
    # Saving the extracted code blocks inside a JSON file
    with open(output_json_path, 'w', encoding='utf-8') as f:
        json.dump(code_blocks, f, ensure_ascii=False, indent=2)

    print(f"Modified Document saved in {output_docx_path}")
    print(f"Extracted code blocks saved in {output_json_path}")

if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: python replace_code_blocks.py <input.docx> <output.docx> <output_codeblocks.json>")
        sys.exit(1)

    main(sys.argv[1], sys.argv[2], sys.argv[3])
